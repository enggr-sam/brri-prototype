#!/usr/bin/env python3
"""Generate BRRI Win2024 field data collection form (source: BRRI PDF only)."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parent / "BRRI_Win2024_Data_Collection_Form.docx"
DOCS_DIR = OUT.parent
REPO_ROOT = DOCS_DIR.parent


def _photo_links_csv_path() -> Path | None:
    for candidate in (
        DOCS_DIR / "photo_folder_links.csv",
        REPO_ROOT / "photo_folder_links.csv",
    ):
        if candidate.is_file():
            return candidate
    return None

PDF_TITLE = (
    "Technical Drawings of BRRI Winnower (Model: BRRI Win2024), "
    "BRRI Publication No. 461, March 2026 — Dr. AKM Saiful Islam"
)

# One example row per collector table; remaining rows are blank templates.
DRIVE_ROOT = "https://drive.google.com/drive/folders/1BRRI_Win2024_Patuakhali_Demo"

PROBLEM_EXAMPLE = (
    "fault.belt",
    "POWER PULLEY BELT",
    "ভি-বেল্ট",
    "Slip / crack / wrong tension",
    "বেল্ট পিছলে যায়",
    "১. মোটর মাউন্ট বোল্ট দিয়ে বেল্ট টেনশন ঠিক করুন। ২. বেল্ট cracked হলে B65 (65 inch) বসান।",
    "10",
)

PROBLEM_BLANKS = [
    ("fault.air", "AIR CONTROL PLATE", "Too much wind — grain loss with chaff"),
    ("fault.sieve_motion", "SIEVE / SIEVE SHAFT", "Sieve not oscillating"),
    ("fault.motor", "MOTOR", "Motor not running / low speed"),
    ("fault.bearing", "PILLOW BEARING BLOCK-206 / BALL BEARING-6302", "Noise / vibration / heat"),
    ("fault.blower", "BLOWER UNITE", "Weak air blast"),
    ("fault.feed", "GRAIN CONTROL PLATE", "Uneven or jammed grain feed"),
    ("fault.multicrop", "SIEVE (THREE TYPE)", "Wrong cleaning — change sieve net"),
]

DEALER_EXAMPLE = (
    "dealer.01",
    "এসি আই মোটরস",
    "হরিংঘাটা মেশিনারীজ, পটুয়াখালী",
    "০১৭১৮২৩২৪০৬",
    "B65 V-belt, UCP206, 6203",
    "B65 মার্কিংসহ বেল্ট রাখে",
)

PHOTO_EXAMPLE = (
    "10",
    "POWER PULLEY BELT (B65 marking visible)",
    "বি-৬৫ বেল্ট",
    "10_POWER_PULLEY_BELT",
    "10_b65_marking.jpg",
    f"{DRIVE_ROOT}/10_POWER_PULLEY_BELT/10_b65_marking.jpg",
    "বেল্ট পিছলে যায়",
)

PHOTO_SLOTS = [
    ("photo.01", "01", "MAIN FRAME", "01_MAIN_FRAME", "01_main_frame.jpg"),
    ("photo.02", "02", "HOPPER (FRONT / BACK / BOTTOM PLATE)", "02_HOPPER", "02_hopper.jpg"),
    ("photo.03", "03", "AIR CONTROL PLATE", "03_AIR_CONTROL_PLATE", "03_air_control_plate.jpg"),
    ("photo.04", "04", "GRAIN CONTROL PLATE", "04_GRAIN_CONTROL_PLATE", "04_grain_control_plate.jpg"),
    ("photo.05", "05", "BLOWER UNITE (assembled)", "05_BLOWER_UNITE", "05_blower_unite.jpg"),
    ("photo.06", "06", "BLOWER COVER PLATE / fan opening", "06_BLOWER_COVER", "06_blower_cover.jpg"),
    ("photo.07", "07", "AIR OUTLET CONTROL PLATE", "07_AIR_OUTLET_CONTROL", "07_air_outlet_control.jpg"),
    ("photo.08", "08", "SIEVE (THREE TYPE)", "08_SIEVE", "08_sieve.jpg"),
    ("photo.09", "09", "SIEVE SHAFT", "09_SIEVE_SHAFT", "09_sieve_shaft.jpg"),
    ("photo.10", "10", "POWER PULLEY BELT (B65 marking visible)", "10_POWER_PULLEY_BELT", "10_b65_marking.jpg"),
    ("photo.11", "11", "MOTOR", "11_MOTOR", "11_motor.jpg"),
    ("photo.12", "12", "MOTOR PULLEY", "12_MOTOR_PULLEY", "12_motor_pulley.jpg"),
    ("photo.13", "13", "BLOWER PULLEY", "13_BLOWER_PULLEY", "13_blower_pulley.jpg"),
    ("photo.14", "14", "PILLOW BEARING BLOCK- UCP206", "14_PILLOW_BEARING_UCP206", "14_pillow_bearing.jpg"),
    ("photo.15", "15", "BALL BEARING-6302", "15_BALL_BEARING_6302", "15_ball_bearing.jpg"),
    ("photo.16", "16", "BEARING 6203 (sieve)", "16_BEARING_6203", "16_bearing_6203.jpg"),
    ("photo.17", "17", "GRAIN OUTLET / GRAIN OUTLET 2ND PLATE", "17_GRAIN_OUTLET", "17_grain_outlet.jpg"),
    ("photo.18", "18", "WINNOWER SHOW COVER", "18_WINNOWER_SHOW_COVER", "18_show_cover.jpg"),
    ("photo.19", "19", "Full machine — front view", "19_FULL_MACHINE_FRONT", "19_full_front.jpg"),
    ("photo.20", "20", "Full machine — side view", "20_FULL_MACHINE_SIDE", "20_full_side.jpg"),
]


def load_photo_drive_links() -> tuple[str, dict[str, str]]:
    """Load folder links from Colab CSV (photo_no → folder_link)."""
    csv_path = _photo_links_csv_path()
    if csv_path is None:
        return "", {}

    root = ""
    by_no: dict[str, str] = {}
    with csv_path.open(encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            no = (row.get("photo_no") or "").strip()
            link = (row.get("folder_link") or "").strip()
            if not no or not link:
                continue
            if no.lower() == "root":
                root = link
            else:
                by_no[no.zfill(2)] = link
    return root, by_no


def problem_rows() -> list[list[str]]:
    rows = [list(PROBLEM_EXAMPLE)]
    for key, part, paper_problem in PROBLEM_BLANKS:
        rows.append([key, part, "", paper_problem, "", "", ""])
    return rows


def dealer_rows() -> list[list[str]]:
    rows = [list(DEALER_EXAMPLE)]
    for n in range(2, 7):
        rows.append([f"dealer.{n:02d}", "", "", "", "", ""])
    return rows


def photo_rows() -> list[list[str]]:
    ex_num, ex_part, ex_local, ex_folder, ex_file, _, ex_problem = PHOTO_EXAMPLE
    _, drive_links = load_photo_drive_links()
    rows: list[list[str]] = []
    for key, num, part, folder, filename in PHOTO_SLOTS:
        link = drive_links.get(num, "")
        if num == ex_num:
            rows.append(
                [key, num, ex_part, ex_local, ex_folder, ex_file, link, ex_problem]
            )
        else:
            rows.append([key, num, part, "", folder, filename, link, ""])
    return rows


def shade(cell, color: str = "E7E6E6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def font9(cell) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)


def add_section(doc: Document, section_id: str, title_bn_en: str) -> None:
    doc.add_page_break()
    h = doc.add_heading(f"SECTION:{section_id} — {title_bn_en}", level=1)


def kv_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    """Rows: (field_key, label, value)."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["field_key", "ক্ষেত্র", "মান"]):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i], "D9E2F3")
    for key, label, value in rows:
        row = t.add_row().cells
        row[0].text = key
        row[1].text = label
        row[2].text = value
        for c in row:
            font9(c)


def grid_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        shade(t.rows[0].cells[i], "D9E2F3")
        font9(t.rows[0].cells[i])
    for data in rows:
        row = t.add_row().cells
        for i, val in enumerate(data):
            row[i].text = val
            font9(row[i])


def main() -> None:
    doc = Document()
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(doc.sections[0], margin, Inches(0.65))

    title = doc.add_heading(
        "BRRI Winnower 2024 — তথ্য ও ছবি সংগ্রহ ফর্ম\n(Data & Photo Collection Form)",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(PDF_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(10)

    doc.add_paragraph(
        "উৎস: BRRI প্রকাশনা (Publication No. 461)। "
        "Google Drive-এ আপলোড করে দলের সদস্যদের সাথে শেয়ার করুন। "
        "ছবি আলাদা ফোল্ডারে রাখুন — নিচের ‘ছবি সংগ্রহ’ অংশে নির্দেশনা আছে।"
    )

    # --- Specs ---
    add_section(doc, "specs", "প্রযুক্তিগত বিবরণ (Technical Specifications)")
    kv_table(
        doc,
        [
            ("machine.model", "মেশিন মডেল", "BRRI Win2024"),
            ("machine.name", "মেশিনের নাম", "BRRI Winnower"),
            ("dimensions.overall_mm", "সামগ্রিক মাপ (mm)", "1350 × 835 × 1310"),
            ("weight.g", "ওজন (g)", "97861.75"),
            ("motor.hp", "মোটর — HP", "1.5"),
            ("motor.kw", "মোটর — kW", "1.1"),
            ("motor.rpm", "মোটর — RPM", "1400"),
            ("belt.type", "পাওয়ার পুলি বেল্ট — TYPE", "B"),
            ("belt.length", "পাওয়ার পুলি বেল্ট — LENGTH", "65 inch (B65)"),
            ("motor_pulley.od_mm", "মোটর পুলি — OD (mm)", "125"),
            ("blower.diameter_mm", "ব্লোয়ার কভার — Ø (mm)", "270"),
            ("bearing.sieve", "ঝরনি বিয়ারিং", "6203 × 2"),
            ("bearing.pillow", "পিলো বিয়ারিং ব্লক", "UCP206 × 2"),
            ("bearing.blower", "বল বিয়ারিং (ব্লোয়ার)", "6302 × 8"),
            ("frame.angle_bar", "মূল ফ্রেম — ANGLE", "38 × 38 × 3 mm"),
            ("frame.flat_bar", "মূল ফ্রেম — FLAT BAR", "20 × 4 mm"),
            ("sieve.types", "ঝরনির ধরন (THREE TYPE OF SIEVE)", "Large / Medium / Small holes net"),
        ],
    )

    # --- Sub-assemblies ---
    add_section(doc, "subassemblies", "উপ-Assembly তালিকা (Sub-Assembly Index)")
    grid_table(
        doc,
        ["field_key", "Sub-Assembly", "DRG / Sheet", "বিবরণ"],
        [
            ("subasm.01", "1 — Main body (BOM)", "Sheet 004–005", "MAIN FRAME, HOPPER, AIR CONTROL PLATE…"),
            ("subasm.02", "2 — SIEVE", "Sheet 022–024", "6203 BEARING QTY 2; THREE TYPE OF SIEVE"),
            ("subasm.03", "3 — BLOWER UNITE", "Sheet 026", "12 parts — fan, shaft, bearings"),
            ("subasm.04", "4 — BEARING HOUSE", "Sheet 039", ""),
            ("subasm.05", "5 — BLOWER PULLEY", "Sheet 040", ""),
            ("subasm.06", "6 — POWER PULLEY BELT", "Sheet 041", "TYPE-B, LENGTH-65 inch"),
            ("subasm.07", "7 — MOTOR", "Sheet 042", "HP-1.5 KW-1.1 RPM-1400"),
            ("subasm.08", "8 — MOTOR PULLEY", "Sheet 043", ""),
            ("subasm.09", "9 — SIEVE SHAFT", "Sheet 044", ""),
            ("subasm.10", "10 — WINNOWER SHOW COVER", "Sheet 045", "QUANTITY-02"),
        ],
    )

    # --- Problems ---
    add_section(doc, "problems", "সমস্যা ও সমাধান (Problems & Solutions)")
    doc.add_paragraph("প্রথম সারি উদাহরণ — বাকি সারিতে আপনার তথ্য লিখুন।")
    grid_table(
        doc,
        [
            "field_key",
            "অংশ (পেপার)",
            "স্থানীয় নাম",
            "সমস্যা (পেপার)",
            "স্থানীয় সমস্যার নাম",
            "সমাধান",
            "ছবি নং",
        ],
        problem_rows(),
    )

    # --- Dealers ---
    add_section(doc, "dealers", "যন্ত্রাংশ সরবরাহকারী (Parts Suppliers)")
    doc.add_paragraph("প্রথম সারি উদাহরণ — বাকি সারিতে ডিলারের তথ্য লিখুন।")
    grid_table(
        doc,
        [
            "field_key",
            "দোকান (বাংলা)",
            "ঠিকানা",
            "মোবাইল",
            "যন্ত্রাংশ",
            "নোট",
        ],
        dealer_rows(),
    )

    # --- Photos ---
    add_section(doc, "photos", "ছবি সংগ্রহ (Photo Collection)")
    doc.add_heading("Google Drive ফোল্ডার কাঠামো (Photo Directory)", level=2)
    doc.add_paragraph(
        "১. Google Drive-এ একটি শেয়ারড ফোল্ডার তৈরি করুন:\n"
        "   BRRI_Win2024_Field_Photos\n"
        "২. প্রতিটি ছবির জন্য উপ-ফোল্ডার (পেপারের DRG NAME অনুযায়ী):\n"
        "   01_MAIN_FRAME/\n"
        "   02_MOTOR/\n"
        "   03_POWER_PULLEY_BELT/\n"
        "   … (নিচের টেবিলের ‘ফোল্ডার নাম’ কলাম)\n"
        "৩. ফাইল নাম: NN_description.jpg — এক ফোল্ডারে একাধিক ছবি OK (উদাহরণ: 10_b65_marking.jpg, 10_belt_worn.jpg)\n"
        "৪. ‘ফাইল নাম’ কলামে শুধু মূল/সেরা ছবির নাম; বাকি ছবি Drive ফোল্ডারেই থাকবে\n"
        "৫. ছবি তোলার পর ‘Google Drive লিংক’ কলামে ফোল্ডার লিংক দিন।"
    )

    photo_headers = [
        "field_key",
        "ছবি নং",
        "অংশ (পেপার)",
        "স্থানীয় নাম",
        "ফোল্ডার নাম",
        "ফাইল নাম (মূল ছবি)",
        "Google Drive লিংক",
        "সম্পর্কিত সমস্যা (স্থানীয়)",
    ]
    doc.add_paragraph("প্রথম সারি উদাহরণ — বাকি সারিতে ছবি ও Google Drive লিংক লিখুন।")
    grid_table(doc, photo_headers, photo_rows())

    # --- Collector meta ---
    add_section(doc, "meta", "সংগ্রহকারীর তথ্য (Collector Info)")
    drive_root, _ = load_photo_drive_links()
    kv_table(
        doc,
        [
            ("meta.collector_name", "নাম", ""),
            ("meta.organization", "প্রতিষ্ঠান", ""),
            ("meta.district", "জেলা", ""),
            ("meta.date", "তারিখ", ""),
            ("meta.drive_root_link", "Google Drive মূল ফোল্ডার লিংক", drive_root),
            ("meta.notes", "মন্তব্য", ""),
        ],
    )

    doc.save(OUT)
    csv_path = _photo_links_csv_path()
    if csv_path:
        print(f"Loaded photo links from {csv_path}")
    else:
        print("No photo_folder_links.csv — photo links left blank.")
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
